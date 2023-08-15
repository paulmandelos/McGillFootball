import pandas as pd
import docx
x = pd.read_excel('/Users/paulmandelos/Desktop/RedBirdsFootball/PlaylistData_2023-04-29.xlsx')

def groupDistances(distance):
    if (distance<=3):
        return 1
    if (distance==5):
        return 3
    if (4<=distance<=7):
        return 2
    if (distance==10):
        return 5
    if (distance>=8):
        return 4
    
def groupLOS(los):
    if (-5<=los<=0):
        return 0
    if (-20<=los<=0):
        return 1
    if (-55<=los<=-20):
        return 2
    if (20<=los<=55):
        return 3
    if (0<=los<=5):
        return 5
    if (0<=los<=20):
        return 4   
    
def WSplay(playname):
    try:
        if "WK" in playname:
            return "WK"
        if "ST" in playname:
            return "ST"
    except:
        return ""
    
x['DISTGROUPED'] = x['DIST'].apply(groupDistances)
x['LNGROUPED'] = x['YARD LN'].apply(groupLOS)
x['SideRun'] = x['OFF PLAY'].apply(WSplay)
groupedDD = x.groupby(by=['DN','DISTGROUPED']).size().to_frame('size')
groupedDDPlayType = x.groupby(by=['DN','DISTGROUPED','PLAY TYPE']).size().to_frame('size')
groupedDDOffensivePlay = x.groupby(by=['DN','DISTGROUPED','OFF PLAY']).size().to_frame('size')

groupedLOS = x.groupby(by=['LNGROUPED']).size().to_frame('size')
groupedLOSRun = x[x['PLAY TYPE']=='Run'].groupby(by=['LNGROUPED']).size().to_frame('size')

groupedLOSPlayType = x.groupby(by=['LNGROUPED','PLAY TYPE']).size().to_frame('size')
groupedLOSOffensivePlay = x.groupby(by=['LNGROUPED','OFF PLAY']).size().to_frame('size')
groupedLOSSide = x[x['PLAY TYPE']=='Run'].groupby(by=['LNGROUPED','SideRun']).size().to_frame('size')

DDPlayType = groupedDDPlayType/groupedDD
DDOffensivePlay = groupedDDOffensivePlay/groupedDD

LOSPlayType = groupedLOSPlayType/groupedLOS
LOSOffensivePlay = groupedLOSOffensivePlay/groupedLOS
LOSSide = groupedLOSSide/groupedLOSRun

DDOffensivePlay.reset_index(inplace=True)
DDOffensivePlay = DDOffensivePlay.groupby(by=['DN','DISTGROUPED']).apply(lambda x: x.sort_values('size', ascending = False).head(3).reset_index(drop = True))
DDOffensivePlay = DDOffensivePlay.drop(columns=['DN','DISTGROUPED'])
DDOffensivePlay.reset_index(inplace=True)

LOSOffensivePlay.reset_index(inplace=True)
LOSOffensivePlay = LOSOffensivePlay.groupby(by=['LNGROUPED']).apply(lambda x: x.sort_values('size', ascending = False).head(3).reset_index(drop = True))
LOSOffensivePlay = LOSOffensivePlay.drop(columns=['LNGROUPED'])
LOSOffensivePlay.reset_index(inplace=True)

LOSSide.reset_index(inplace=True)
LOSPlayType.reset_index(inplace=True)

###Creating the Word document to output the things
DDPlayType.reset_index(inplace=True)
print(LOSOffensivePlay)
print(LOSPlayType)
print(LOSSide)
document = docx.Document()
document.add_heading('Montreal Tendencies', 0)
document.add_paragraph("STATS FOR DOWN AND DISTANCE\n\n")
table = document.add_table(rows=18, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '0 & 10'
hdr_cells[1].text = '1 & 10'
hdr_cells1 = table.rows[1].cells
try:
    hdr_cells1[0].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==0.0) & (DDOffensivePlay['DISTGROUPED']==5) & (DDOffensivePlay['level_2']==0)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==0.0) & (DDOffensivePlay['DISTGROUPED']==5) & (DDOffensivePlay['level_2']==0)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells1[0].text = "--"
try:
    hdr_cells1[1].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==1.0) & (DDOffensivePlay['DISTGROUPED']==5) & (DDOffensivePlay['level_2']==0)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==1.0) & (DDOffensivePlay['DISTGROUPED']==5) & (DDOffensivePlay['level_2']==0)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells1[1].text = "--"

hdr_cells2 = table.rows[2].cells
try:
    hdr_cells2[0].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==0.0) & (DDOffensivePlay['DISTGROUPED']==5) & (DDOffensivePlay['level_2']==1)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==0.0) & (DDOffensivePlay['DISTGROUPED']==5) & (DDOffensivePlay['level_2']==1)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells2[0].text = "--"
try:
    hdr_cells2[1].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==1.0) & (DDOffensivePlay['DISTGROUPED']==5) & (DDOffensivePlay['level_2']==1)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==1.0) & (DDOffensivePlay['DISTGROUPED']==5) & (DDOffensivePlay['level_2']==1)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells2[1].text = "--"

hdr_cells3 = table.rows[3].cells
try:
    hdr_cells3[0].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==0.0) & (DDOffensivePlay['DISTGROUPED']==5) & (DDOffensivePlay['level_2']==2)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==0.0) & (DDOffensivePlay['DISTGROUPED']==5) & (DDOffensivePlay['level_2']==2)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells3[0].text = "--"
try:
    hdr_cells3[1].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==1.0) & (DDOffensivePlay['DISTGROUPED']==5) & (DDOffensivePlay['level_2']==2)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==1.0) & (DDOffensivePlay['DISTGROUPED']==5) & (DDOffensivePlay['level_2']==2)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells3[1].text = "--"

hdr_cells4 = table.rows[4].cells
try:
    hdr_cells4[0].text = str(round(DDPlayType[(DDPlayType['DN']==0.0) & (DDPlayType['DISTGROUPED']==5.0) & (DDPlayType['PLAY TYPE']=='Pass')]['size'].values[0]*100,1)) + '% Pass'
except:
    hdr_cells4[0].text = "--"
try:
    hdr_cells4[1].text = str(round(DDPlayType[(DDPlayType['DN']==1.0) & (DDPlayType['DISTGROUPED']==5.0) & (DDPlayType['PLAY TYPE']=='Pass')]['size'].values[0]*100,1)) + '% Pass'
except:
    hdr_cells4[1].text = "--"
hdr_cells5 = table.rows[5].cells
try:
    hdr_cells5[0].text = str(round(DDPlayType[(DDPlayType['DN']==0.0) & (DDPlayType['DISTGROUPED']==5.0) & (DDPlayType['PLAY TYPE']=='Run')]['size'].values[0]*100,1)) + '% Run'
except:
    hdr_cells5[0].text = "--"
try:
    hdr_cells5[1].text = str(round(DDPlayType[(DDPlayType['DN']==1.0) & (DDPlayType['DISTGROUPED']==5.0) & (DDPlayType['PLAY TYPE']=='Run')]['size'].values[0]*100,1)) + '% Run'
except:
    hdr_cells5[1].text = "--"
hdr_cells6 = table.rows[6].cells
hdr_cells6[0].text = '1 & 5'
hdr_cells6[1].text = '2 & 8+'
hdr_cells7 = table.rows[7].cells
try:
    hdr_cells7[0].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==1.0) & (DDOffensivePlay['DISTGROUPED']==3) & (DDOffensivePlay['level_2']==0)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==1.0) & (DDOffensivePlay['DISTGROUPED']==3) & (DDOffensivePlay['level_2']==0)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells7[0].text = "--"
try:
    hdr_cells7[1].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==4) & (DDOffensivePlay['level_2']==0)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==4) & (DDOffensivePlay['level_2']==0)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells7[1].text = "--"
hdr_cells8 = table.rows[8].cells
try:
    hdr_cells8[0].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==1.0) & (DDOffensivePlay['DISTGROUPED']==3) & (DDOffensivePlay['level_2']==1)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==1.0) & (DDOffensivePlay['DISTGROUPED']==3) & (DDOffensivePlay['level_2']==1)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells8[0].text = "--"
try:
    hdr_cells8[1].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==4) & (DDOffensivePlay['level_2']==1)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==4) & (DDOffensivePlay['level_2']==1)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells8[1].text = "--"
hdr_cells9 = table.rows[9].cells
try:
    hdr_cells9[0].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==1.0) & (DDOffensivePlay['DISTGROUPED']==3) & (DDOffensivePlay['level_2']==2)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==1.0) & (DDOffensivePlay['DISTGROUPED']==3) & (DDOffensivePlay['level_2']==2)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells9[0].text = "--"
try:
    hdr_cells9[1].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==4) & (DDOffensivePlay['level_2']==2)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==4) & (DDOffensivePlay['level_2']==2)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells9[1].text = "--"
hdr_cells10 = table.rows[10].cells
try:
    hdr_cells10[0].text = str(round(DDPlayType[(DDPlayType['DN']==1.0) & (DDPlayType['DISTGROUPED']==2) & (DDPlayType['PLAY TYPE']=='Pass')]['size'].values[0]*100,1)) + '% Pass'
except:
    hdr_cells10[0].text = "--"
try:
    hdr_cells10[1].text = str(round(DDPlayType[(DDPlayType['DN']==2.0) & (DDPlayType['DISTGROUPED']==2) & (DDPlayType['PLAY TYPE']=='Pass')]['size'].values[0]*100,1)) + '% Pass'
except:
    hdr_cells10[1].text = "--"
hdr_cells11 = table.rows[11].cells
try:
    hdr_cells11[0].text = str(round(DDPlayType[(DDPlayType['DN']==1.0) & (DDPlayType['DISTGROUPED']==2) & (DDPlayType['PLAY TYPE']=='Run')]['size'].values[0]*100,1)) + '% Run'
except:
    hdr_cells11[0].text = "--"
try:
    hdr_cells11[1].text = str(round(DDPlayType[(DDPlayType['DN']==2.0) & (DDPlayType['DISTGROUPED']==2) & (DDPlayType['PLAY TYPE']=='Run')]['size'].values[0]*100,1)) + '% Run'
except:
    hdr_cells11[1].text = "--"

hdr_cells12 = table.rows[12].cells
hdr_cells12[0].text = '2 & 4-7'
hdr_cells12[1].text = '2 & 3-'
hdr_cells13 = table.rows[13].cells
try:
    hdr_cells13[0].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==2) & (DDOffensivePlay['level_2']==0)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==2) & (DDOffensivePlay['level_2']==0)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells13[0].text = "--"
try:
    hdr_cells13[1].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==1) & (DDOffensivePlay['level_2']==0)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==1) & (DDOffensivePlay['level_2']==0)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells13[1].text = "--"
hdr_cells14 = table.rows[14].cells
try:
    hdr_cells14[0].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==2) & (DDOffensivePlay['level_2']==1)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==2) & (DDOffensivePlay['level_2']==1)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells14[0].text = "--"
try:
    hdr_cells14[1].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==1) & (DDOffensivePlay['level_2']==1)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==1) & (DDOffensivePlay['level_2']==1)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells14[1].text = "--"
hdr_cells15 = table.rows[15].cells
try:
    hdr_cells15[0].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==2) & (DDOffensivePlay['level_2']==2)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==2) & (DDOffensivePlay['level_2']==2)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells15[0].text = "--"
try:
    hdr_cells15[1].text = str(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==1) & (DDOffensivePlay['level_2']==2)]['OFF PLAY'].values[0])+" : "+str(round(DDOffensivePlay[(DDOffensivePlay['DN']==2.0) & (DDOffensivePlay['DISTGROUPED']==1) & (DDOffensivePlay['level_2']==2)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells15[1].text = "--"
hdr_cells16 = table.rows[16].cells
try:
    hdr_cells16[0].text = str(round(DDPlayType[(DDPlayType['DN']==2.0) & (DDPlayType['DISTGROUPED']==2) & (DDPlayType['PLAY TYPE']=='Pass')]['size'].values[0]*100,1)) + '% Pass'
except:
    hdr_cells16[0].text = "--"
try:
    hdr_cells16[1].text = str(round(DDPlayType[(DDPlayType['DN']==2.0) & (DDPlayType['DISTGROUPED']==1) & (DDPlayType['PLAY TYPE']=='Pass')]['size'].values[0]*100,1)) + '% Pass'
except:
    hdr_cells16[1].text = "--"
hdr_cells17 = table.rows[17].cells
try:
    hdr_cells17[0].text = str(round(DDPlayType[(DDPlayType['DN']==2.0) & (DDPlayType['DISTGROUPED']==2) & (DDPlayType['PLAY TYPE']=='Run')]['size'].values[0]*100,1)) + '% Run'
except:
    hdr_cells17[0].text = "--"
try:
    hdr_cells17[1].text = str(round(DDPlayType[(DDPlayType['DN']==2.0) & (DDPlayType['DISTGROUPED']==1) & (DDPlayType['PLAY TYPE']=='Run')]['size'].values[0]*100,1)) + '% Run'
except:
    hdr_cells17[1].text = "--"

document.add_page_break()
document.add_paragraph("STATS FOR Area on Field\n\n")
table2 = document.add_table(rows=18, cols=2)
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Inside Own 5'
hdr_cells[1].text = 'Inside Own 20'
hdr_cells1 = table2.rows[1].cells
try:
    hdr_cells1[0].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==0.0) & (LOSOffensivePlay['level_1']==0)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==0.0) & (LOSOffensivePlay['level_1']==0)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells1[0].text = "--"
try:
    hdr_cells1[1].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==1.0) & (LOSOffensivePlay['level_1']==0)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==1.0) & (LOSOffensivePlay['level_1']==0)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells1[1].text = "--"

hdr_cells2 = table2.rows[2].cells
try:
    hdr_cells2[0].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==0.0) & (LOSOffensivePlay['level_1']==1)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==0.0) & (LOSOffensivePlay['level_1']==1)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells2[0].text = "--"
try:
    hdr_cells2[1].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==1.0) & (LOSOffensivePlay['level_1']==1)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==1.0) & (LOSOffensivePlay['level_1']==1)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells2[1].text = "--"

hdr_cells3 = table2.rows[3].cells
try:
    hdr_cells3[0].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==0.0) & (LOSOffensivePlay['level_1']==2)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==0.0) & (LOSOffensivePlay['level_1']==2)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells3[0].text = "--"
try:
    hdr_cells3[1].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==1.0) & (LOSOffensivePlay['level_1']==2)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==1.0) & (LOSOffensivePlay['level_1']==2)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells3[1].text = "--"

hdr_cells4 = table2.rows[4].cells
try:
    hdr_cells4[0].text = str(round(LOSPlayType[(LOSPlayType['LNGROUPED']==0.0)  & (LOSPlayType['PLAY TYPE']=='Pass')]['size'].values[0]*100,1)) + '% Pass'
except:
    hdr_cells4[0].text = "--"
try:
    hdr_cells4[1].text = str(round(LOSPlayType[(LOSPlayType['LNGROUPED']==1.0)  & (LOSPlayType['PLAY TYPE']=='Pass')]['size'].values[0]*100,1)) + '% Pass'
except:
    hdr_cells4[1].text = "--"
hdr_cells5 = table2.rows[5].cells
try:
    hdr_cells5[0].text = str(round(LOSPlayType[(LOSPlayType['LNGROUPED']==0.0)  & (LOSPlayType['PLAY TYPE']=='Run')]['size'].values[0]*100,1)) + '% Run'
except:
    hdr_cells5[0].text = "--"
try:
    hdr_cells5[1].text = str(round(LOSPlayType[(LOSPlayType['LNGROUPED']==1.0)  & (LOSPlayType['PLAY TYPE']=='Run')]['size'].values[0]*100,1)) + '% Run'
except:
    hdr_cells5[1].text = "--"
hdr_cells6 = table2.rows[6].cells
hdr_cells6[0].text = 'Own 20 to Midfield'
hdr_cells6[1].text = 'Opp 20 to Midfield'
hdr_cells7 = table2.rows[7].cells
try:
    hdr_cells7[0].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==1.0) & (LOSOffensivePlay['level_1']==0)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==1.0) & (LOSOffensivePlay['level_1']==0)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells7[0].text = "--"
try:
    hdr_cells7[1].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==0)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==0)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells7[1].text = "--"
hdr_cells8 = table2.rows[8].cells
try:
    hdr_cells8[0].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==1.0) & (LOSOffensivePlay['level_1']==1)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==1.0) & (LOSOffensivePlay['level_1']==1)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells8[0].text = "--"
try:
    hdr_cells8[1].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==1)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==1)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells8[1].text = "--"
hdr_cells9 = table2.rows[9].cells
try:
    hdr_cells9[0].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==1.0) & (LOSOffensivePlay['level_1']==2)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==1.0) & (LOSOffensivePlay['level_1']==2)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells9[0].text = "--"
try:
    hdr_cells9[1].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==2)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==2)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells9[1].text = "--"
hdr_cells10 = table2.rows[10].cells
try:
    hdr_cells10[0].text = str(round(LOSPlayType[(LOSPlayType['LNGROUPED']==1.0)  & (LOSPlayType['PLAY TYPE']=='Pass')]['size'].values[0]*100,1)) + '% Pass'
except:
    hdr_cells10[0].text = "--"
try:
    hdr_cells10[1].text = str(round(LOSPlayType[(LOSPlayType['LNGROUPED']==2.0)  & (LOSPlayType['PLAY TYPE']=='Pass')]['size'].values[0]*100,1)) + '% Pass'
except:
    hdr_cells10[1].text = "--"
hdr_cells11 = table2.rows[11].cells
try:
    hdr_cells11[0].text = str(round(LOSPlayType[(LOSPlayType['LNGROUPED']==1.0)  & (LOSPlayType['PLAY TYPE']=='Run')]['size'].values[0]*100,1)) + '% Run'
except:
    hdr_cells11[0].text = "--"
try:
    hdr_cells11[1].text = str(round(LOSPlayType[(LOSPlayType['LNGROUPED']==2.0)  & (LOSPlayType['PLAY TYPE']=='Run')]['size'].values[0]*100,1)) + '% Run'
except:
    hdr_cells11[1].text = "--"

hdr_cells12 = table2.rows[12].cells
hdr_cells12[0].text = 'Inside Opp 20'
hdr_cells12[1].text = 'Inside Opp 5'
hdr_cells13 = table2.rows[13].cells
try:
    hdr_cells13[0].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==0)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==0)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells13[0].text = "--"
try:
    hdr_cells13[1].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==0)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==0)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells13[1].text = "--"
hdr_cells14 = table2.rows[14].cells
try:
    hdr_cells14[0].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==1)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==1)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells14[0].text = "--"
try:
    hdr_cells14[1].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==1)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==1)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells14[1].text = "--"
hdr_cells15 = table2.rows[15].cells
try:
    hdr_cells15[0].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==2)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==2)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells15[0].text = "--"
try:
    hdr_cells15[1].text = str(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==2)]['OFF PLAY'].values[0])+" : "+str(round(LOSOffensivePlay[(LOSOffensivePlay['LNGROUPED']==2.0) & (LOSOffensivePlay['level_1']==2)]['size'].values[0]*100,2))+'%'
except:
    hdr_cells15[1].text = "--"
hdr_cells16 = table2.rows[16].cells
try:
    hdr_cells16[0].text = str(round(LOSPlayType[(LOSPlayType['LNGROUPED']==2.0)  & (LOSPlayType['PLAY TYPE']=='Pass')]['size'].values[0]*100,1)) + '% Pass'
except:
    hdr_cells16[0].text = "--"
try:
    hdr_cells16[1].text = str(round(LOSPlayType[(LOSPlayType['LNGROUPED']==2.0)  & (LOSPlayType['PLAY TYPE']=='Pass')]['size'].values[0]*100,1)) + '% Pass'
except:
    hdr_cells16[1].text = "--"
hdr_cells17 = table2.rows[17].cells
try:
    hdr_cells17[0].text = str(round(LOSPlayType[(LOSPlayType['LNGROUPED']==2.0)  & (LOSPlayType['PLAY TYPE']=='Run')]['size'].values[0]*100,1)) + '% Run'
except:
    hdr_cells17[0].text = "--"
try:
    hdr_cells17[1].text = str(round(LOSPlayType[(LOSPlayType['LNGROUPED']==2.0)  & (LOSPlayType['PLAY TYPE']=='Run')]['size'].values[0]*100,1)) + '% Run'
except:
    hdr_cells17[1].text = "--"

document.save('/Users/paulmandelos/Desktop/RedBirdsFootball/test.docx')
